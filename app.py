import streamlit as st
import moviepy.editor as mp
import speech_recognition as sr
import os
import nltk
import tempfile
from gensim.summarization import summarize
from PyPDF2 import PdfReader
from pattern.web import Google
from pattern.en import pluralize, singularize, comparative, superlative
from googletrans import Translator
import codecs
from pytube import YouTube
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate


# Load environment variables
load_dotenv()
genai.configure 

# Download NLTK resources
nltk.download('punkt')
nltk.download('wordnet')
nltk.download('stopwords')
from nltk.corpus import wordnet
from textblob import TextBlob


# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as f:
        reader = PdfReader(f)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

# Function to summarize text using gensim
def summary_gensim(text):
    return summarize(text)

# Function to transcribe video
def transcribe_video(video_path):
    clip = mp.VideoFileClip(video_path)
    recognizer = sr.Recognizer()
    text = ""
    for i, segment in enumerate(range(0, int(clip.duration), 30)):
        audio_path = f"temp_audio_{i}.wav"
        clip.subclip(segment, min(segment + 30, int(clip.duration))).audio.write_audiofile(audio_path)
        with sr.AudioFile(audio_path) as source:
            audio_data = recognizer.record(source)
            segment_text = recognizer.recognize_google(audio_data)
            text += segment_text + " "
        os.remove(audio_path)
    clip.close()
    return text

def download_youtube_video(url):
    try:
        yt = YouTube(url)
        video = yt.streams.get_highest_resolution()
        video_path = "temp_video.mp4"
        video.download() 
        return video_path
    except Exception as e:
        st.error(f"Error downloading YouTube video: {e}")
        return None
    
def save_summary_as_word(summary_text, file_name="summary.docx"):
    doc = docx.Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(summary_text)
    doc.save(file_name)
    return file_name

def get_pdf_text(pdf_docs):
    text=""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

def user_input(user_question):
    if not os.path.exists("faiss_index/index.faiss"):
        st.error("Index file not found. Please create the index first by processing the PDF.")
        return

    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()

    response = chain(
        {"input_documents": docs, "question": user_question},
        return_only_outputs=True
    )

    st.write("Reply: ", response["output_text"])

def main():
    
    st.set_page_config(layout="wide")
    st.title("SummarEase")
    # Add a "Home" button at the top left
    st.sidebar.markdown(
        """
        <a href="" style="text-decoration: none; color: black; font-size: 16px;">
            🏠 Home
        </a>
        <hr style="border: none; border-top: 1px solid #ccc; margin-top: 10px; margin-bottom: 10px;">
        """, 
        unsafe_allow_html=True
    )

    st.sidebar.write("Select Activity of your choice below: ")
    st.sidebar.write("- Summarize Text")
    st.sidebar.write("- Summarize Document")
    st.sidebar.write("- Ask Your PDF")
    st.sidebar.write("- Video Transcription")
    st.sidebar.write("- Synonyms")
    st.sidebar.write("- Translator")
    st.sidebar.write("- Search")
    st.sidebar.write("- Spell Correction")
    


    st.sidebar.write("Select Activity of your choice below:")
    activities = ["Summarize Text", "Summarize Document","Ask Your PDF", "Video Transcription", "Synonyms", "Translator", "Search", "Spell Correction"]
    choice = st.sidebar.selectbox("Select Activity", activities)

    if choice == "Summarize Text":
        st.subheader("Quickly condense long text into key points.")
        text_range = st.sidebar.slider("Summarize words Range", 25, 500)
        text = st.text_area("Input Text For Summary", height=250)
        if st.button("Summarize based on range"):
            summary = summary_gensim(text)
            st.warning(summary)

    elif choice == "Video Transcription":
        st.subheader("Convert video audio to text effortlessly.")
        uploaded_file = st.file_uploader("Upload a video file", type=["mp4", "avi"])
        if uploaded_file is not None:
            st.video(uploaded_file)
            if st.button("Transcribe Video"):
                temp_dir = tempfile.mkdtemp()
                temp_video_path = os.path.join(temp_dir, "temp_video.mp4")
                with open(temp_video_path, "wb") as f:
                    f.write(uploaded_file.read())
                text = transcribe_video(temp_video_path)
                st.subheader("Transcription:")
                st.write(text)
                os.remove(temp_video_path)
                os.rmdir(temp_dir)

    elif choice == "Summarize Document":
        st.subheader("Generate summaries for uploaded documents.")
        input_file = st.file_uploader("Upload your document here", type=['pdf'])
        if input_file is not None:
            if st.button("Summarize Document"):
                with open("doc_file.pdf", "wb") as f:
                    f.write(input_file.getbuffer())
                col1, col2 = st.columns([1, 1])
                with col1:
                    st.info("File uploaded successfully")
                    extracted_text = extract_text_from_pdf("doc_file.pdf")
                    st.markdown("**Extracted Text is Below:**")
                    st.info(extracted_text)
                with col2:
                    st.markdown("**Summary Result**")
                    text = extract_text_from_pdf("doc_file.pdf")
                    doc_summary = summary_gensim(text)
                    st.success(doc_summary)

                # Create the Word file from the summary
                word_file = save_summary_as_word(doc_summary)

                # Create a "Download Summary" button for Word file
                with open(word_file, "rb") as f:
                    st.download_button(
                        label="Download Summary",
                        data=f,
                        file_name='summary.docx',
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )

    elif choice == "Synonyms":
        st.subheader("Find alternative words and expand your vocabulary.")
        text = st.text_area("Enter Text")
        if st.button("Find"):
            for syn in nltk.corpus.wordnet.synsets(text):
                for i in syn.lemmas():
                    st.success(i.name())
        if st.checkbox("Definition"):
            for syn in nltk.corpus.wordnet.synsets(text):
                st.warning(syn.definition())
        if st.checkbox("Example"):
            for syn in nltk.corpus.wordnet.synsets(text):
                st.success(syn.examples())

    elif choice == "Translator":
        st.subheader("Translate text into multiple languages instantly.")
        row_text = st.text_area("Enter Your Text For Translation", height=300)
        selected_language = st.selectbox("Select Language", ["english", "tamil", "punjabi", "gujarati", "hindi", "urdu", "kannada", "bengali", "telugu", "marathi"])
        translator = Translator()
        if st.button("Translate"):
            translated_text = translator.translate(row_text, dest=selected_language).text
            st.success(translated_text)

    elif choice == "Search":
        st.subheader("Discover relevant content with smart search.")
        row_text = st.text_input("Search Anything")
        google = Google(license=None)
        if st.button("Search"):
            for search_result in google.search(row_text):
                st.write(search_result.text)
                st.warning(search_result.url)

    elif choice == "Spell Correction":
        st.subheader("Improve spelling accuracy in your text.")
        text_data = st.text_area("Enter Text Here")
        a = TextBlob(text_data)
        if st.button("Correct"):
            st.success(a.correct())
        st.subheader("Pluralize & Singularize")
        text_data1 = st.text_input("Enter a word For Pluralize/Singularize")
        if st.checkbox("Pluralize"):
            st.warning(pluralize(text_data1))
        if st.checkbox("Singularize"):
            st.warning(singularize(text_data1))
        st.subheader("Comparative & Superlative")
        if st.checkbox("Comparative"):
            st.warning(comparative(text_data1))
        if st.checkbox("Superlative"):
            st.warning(superlative(text_data1))

    elif choice == "Download Lecture":
        st.subheader("Download Lecture from YouTube")
        url = st.text_input("Enter YouTube Video URL:")
        if st.button("Download"):
            video_path = download_youtube_video(url)
            if video_path:
                st.success("Video Downloaded Successfully!")
                st.video(video_path)

if __name__ == "__main__":
    main()
