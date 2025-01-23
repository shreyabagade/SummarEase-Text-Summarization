import streamlit as st
import moviepy.editor as mp
import speech_recognition as sr
import os
import nltk
import tempfile
from gensim.summarization import summarize
from PyPDF2 import PdfReader
from pattern.web import Google
from pattern.en import pluralize, singularize, comparative, superlative
from googletrans import Translator
import codecs
from pytube import YouTube
import docx

# Download NLTK resources
nltk.download('punkt')
nltk.download('wordnet')
nltk.download('stopwords')
from nltk.corpus import wordnet
from textblob import TextBlob

# Load spaCy model
import spacy
nlp = spacy.load('en_core_web_sm')

# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    with open(file_path, "rb") as f:
        reader = PdfReader(f)
        text = ""
        
        # Loop through all the pages
        for page in reader.pages:
            text += page.extract_text() + "\n"  # Extract text and append a newline for clarity
            
    return text


# Custom Components Function
def st_calculator(calc_html,width=1000,height=1350):
    calc_file = codecs.open(calc_html,'r')
    page = calc_file.read()
    components.html(page,width=width,height=height,scrolling=False)

# Function to summarize text using gensim
def summary_gensim(text):
    return summarize(text)

# Function to transcribe video
def transcribe_video(video_path):
    # Load video clip
    clip = mp.VideoFileClip(video_path)

    # Transcribe audio for each 30-second segment
    recognizer = sr.Recognizer()
    text = ""
    for i, segment in enumerate(range(0, int(clip.duration), 30)):
        audio_path = f"temp_audio_{i}.wav"
        clip.subclip(segment, min(segment + 30, int(clip.duration))).audio.write_audiofile(audio_path)

        with sr.AudioFile(audio_path) as source:
            audio_data = recognizer.record(source)
            segment_text = recognizer.recognize_google(audio_data)
            text += segment_text + " "

        # Clean up: remove the temporary audio file
        os.remove(audio_path)

    # Close the video clip
    clip.close()

    return text

def download_youtube_video(url):
    try:
        yt = YouTube(url)
        video = yt.streams.get_highest_resolution()
        video_path = "temp_video.mp4"  # Change this to a relative path
        video.download()  # Removed the video_path argument since it's unnecessary
        return video_path
    except Exception as e:
        st.error(f"Error downloading YouTube video: {e}")
        return None
    
def save_summary_as_word(summary_text, file_name="summary.docx"):
    doc = docx.Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(summary_text)
    # Save the file
    doc.save(file_name)
    # Return the file name to be used for download
    return file_name

def main():
    st.set_page_config(layout="wide")
    st.title("SummarEase")

    st.sidebar.write("Select Activity of your choice below: ")
    st.sidebar.write("- Summarize Text and Transcribe Video")
    st.sidebar.write("- Summarize Document")
    st.sidebar.write("- Synonyms")
    st.sidebar.write("- Translator")
    st.sidebar.write("- Search")
    st.sidebar.write("- Spell Correction")

    activities = ["Summarize Text", "Summarize Document", "Synonyms", "Translator", "Search", "Spell Correction", "Download Lecture"]
    choice = st.sidebar.selectbox("Select Activity", activities)

    if choice == "Summarize Text":
        text_range = st.sidebar.slider("Summarize words Range", 25, 500)
        text = st.text_area("Input Text For Summary", height=250)
        if st.button("Summarize based on range"):
            summary = summary_gensim(text, word_count=text_range)
            st.warning(summary)
        
        # Video Transcription
        st.title("Video Transcription")
        uploaded_file = st.file_uploader("Upload a video file", type=["mp4", "avi"])
        if uploaded_file is not None:
            st.video(uploaded_file)

            if st.button("Transcribe Video"):
                # Create a temporary directory
                temp_dir = tempfile.mkdtemp()

                # Write the file to the temporary directory
                temp_video_path = os.path.join(temp_dir, "temp_video.mp4")
                with open(temp_video_path, "wb") as f:
                    f.write(uploaded_file.read())

                text = transcribe_video(temp_video_path)
                st.subheader("Transcription:")
                st.write(text)

                # Clean up: remove the temporary directory and its contents
                os.remove(temp_video_path)
                os.rmdir(temp_dir)

    elif choice == "Summarize Document":
        st.subheader("Summarize Document from PDF")
        input_file = st.file_uploader("Upload your document here", type=['pdf'])
        if input_file is not None:
            if st.button("Summarize Document"):
                with open("doc_file.pdf", "wb") as f:
                    f.write(input_file.getbuffer())
                col1, col2 = st.columns([1, 1])
                with col1:
                    st.info("File uploaded successfully")
                    extracted_text = extract_text_from_pdf("doc_file.pdf")
                    st.markdown("**Extracted Text is Below:**")
                    st.info(extracted_text)
                with col2:
                    st.markdown("**Summary Result**")
                    text = extract_text_from_pdf("doc_file.pdf")
                    doc_summary = summary_gensim(text)
                    st.success(doc_summary)

                # Create the Word file from the summary
                word_file = save_summary_as_word(doc_summary)

                # Create a "Download Summary" button for Word file
                with open(word_file, "rb") as f:
                    st.download_button(
                        label="Download Summary",
                        data=f,
                        file_name='summary.docx',
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )

    elif choice == "Synonyms":
        st.subheader("Words Synonyms")
        text = st.text_area("Enter Text")
        if st.button("Find"):
            for syn in wordnet.synsets(text):
                for i in syn.lemmas():
                    st.success(i.name())
        if st.checkbox("Definition"):
            for syn in wordnet.synsets(text):
                st.warning(syn.definition())
        if st.checkbox("Example"):
            for syn in wordnet.synsets(text):
                st.success(syn.examples())

    elif choice == "Translator":
        st.subheader("Text Translator")
        row_text = st.text_area("Enter Your Text For Translation", height=300)
        selected_language = st.selectbox("Select Language", ["english", "tamil", "punjabi", "gujrati", "hindi", "urdu", "kannada", "bengali", "telugu","marathi"])
        translator = Translator()
        if st.button("Translate"):
            translated_text = translator.translate(row_text, dest=selected_language).text
            st.success(translated_text)

    elif choice == "Search":
        st.subheader("Search Bar")
        row_text = st.text_input("Search Anything")
        google = Google(license=None)
        if st.button("Search"):
            for search_result in google.search(row_text):
                st.write(search_result.text)
                st.warning(search_result.url)

    elif choice == "Spell Correction":
        st.subheader("Spell Correction")
        text_data = st.text_area("Enter Text Here")
        a = TextBlob(text_data)
        if st.button("Correct"):
            st.success(a.correct())
        st.subheader("Pluralize & singularize")
        text_data1 = st.text_input("Enter a word For pluralize / singularize")
        if st.checkbox("Pluralize"):
            st.warning(pluralize(text_data1))
        if st.checkbox("Singularize"):
            st.warning(singularize(text_data1))
        st.subheader("Comparative & superlative")
        text2 = st.text_input("Enter Text For comparative & superlative")
        if st.checkbox("Comparative"):
            st.success(comparative(text2))
        if st.checkbox("Superlative"):
            st.success(superlative(text2))

    elif choice == "Download Lecture":
        st.title("Download Lecture from Youtube")

        url = st.text_input("Enter YouTube video URL of the Lecture:")

        if url:
            if st.button("Download Video"):
                video_path = download_youtube_video(url)
                if video_path:
                    st.success("Video downloaded successfully!")
            else:
                st.warning("Please enter a YouTube video URL and click 'Download Video'.")

if __name__ == '__main__':
    main()
